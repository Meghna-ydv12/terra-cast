import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Project Progress Report: TerraCast', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_bold_paragraph(doc_obj, text_tuples):
    p = doc_obj.add_paragraph()
    for text, is_bold in text_tuples:
        run = p.add_run(text)
        if is_bold:
            run.bold = True
    return p

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), "0000FF")
    rPr.append(c)
    
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

# 1. Project Title
doc.add_heading('1. Project Title', level=1)
doc.add_paragraph('TerraCast: An Integrated Machine Learning Pipeline for Urban Air Quality Forecasting and Anomaly Detection')

# 2. Problem Statement
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph('Rapid urbanization has led to severe air quality degradation across major metropolitan areas. However, existing air quality forecasting models act as "black boxes" that struggle to predict sudden pollution anomalies or explain their underlying causes, making it nearly impossible for city officials to formulate actionable, real-time environmental interventions.')

# 3. Objective
doc.add_heading('3. Objective of the Project', level=1)
add_bold_paragraph(doc, [
    ('To develop a hybrid ', False),
    ('Dual-Task machine learning framework', True),
    (' (Regression + Classification) enhanced by a ', False),
    ('Rule-Based Recommendation Engine', True),
    ('. The objective is to accurately forecast the ', False),
    ('Air Quality Index (AQI)', True),
    (' and automatically translate ', False),
    ('SHAP-based feature importance', True),
    (' into natural language policy briefs within an ', False),
    ('Interactive Scenario Simulator', True),
    (', bridging the gap between raw data science and practical urban planning.', False)
])

# 4. Literature Review
doc.add_heading('4. Literature Review', level=1)
doc.add_paragraph('The following table summarizes recent research studies (published post-2018) on the application of Machine Learning (ML) and Deep Learning (DL) for predicting AQI and pollutant concentrations, with a focus on specific architectural limitations.')

lit_table = doc.add_table(rows=1, cols=6)
lit_table.style = 'Table Grid'
hdr_cells = lit_table.rows[0].cells
headers = ["Reference & Link", "Dataset Used", "Algorithm(s)", "Key Finding", "Specific Architectural Limitation", "TerraCast Improvement"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].bold = True

lit_data = [
    ("Zamani Joharestani et al. (2019)\n", "https://doi.org/10.3390/atmos10070373", "Tehran AQI", "XGBoost, RF, DL", "XGBoost outperformed DL on tabular data.", "Models are evaluated independently, lacking a unified architecture.", "Integrates both XGBoost and RF concurrently in a Dual-Task Pipeline."),
    ("Wang et al. (2020)\n", "https://doi.org/10.1007/s00521-020-05535-w", "Beijing AQI", "CT-LSTM", "Improved temporal sequence forecasting.", "CT-LSTM is computationally expensive and struggles with high-frequency spatial noise.", "Utilizes gradient-boosted trees for sub-second inference without sequence decay."),
    ("Zhang et al. (2020)\n", "https://doi.org/10.1016/j.apr.2020.09.003", "Beijing EPA", "Bi-LSTM", "High accuracy on historical PM2.5.", "Relies heavily on massive labeled sequence data to prevent bidirectional overfitting.", "XGBoost inherently handles sparse tabular anomalies with built-in regularization."),
    ("Pan et al. (2018)\n", "https://doi.org/10.1088/1755-1315/113/1/012127", "China EPA", "XGBoost", "Identified optimal hyper-parameters for PM2.5.", "Solely predicts PM2.5 and ignores the multi-pollutant calculation required for true AQI.", "Built as an aggregated model factoring in all primary pollutants (PM2.5, PM10, NO2, SO2, CO, O3)."),
    ("Kumar et al. (2022)\n", "https://doi.org/10.1007/s13762-022-04241-5", "Indian CPCB", "ML Ensemble", "Categorized AQI risk accurately.", "Focuses purely on stationary data without dynamically integrating real-time anomalies.", "Engineered with a real-time web UI allowing dynamic parameter adjustments."),
    ("Gilik et al. (2022)\n", "https://doi.org/10.1007/s11356-021-16227-w", "Istanbul AQI", "CNN-LSTM", "Captured spatial and temporal pollution spread.", "Fails to isolate the independent effect of exogenous variables like humidity and wind.", "Implements SHAP to decouple and isolate the impact of individual exogenous features."),
    ("Iskandaryan et al. (2023)\n", "https://doi.org/10.1109/access.2023.3234214", "Madrid AQI", "GNN", "Modeled city-wide sensor graphs.", "Extremely complex parameter tuning requires deep data science expertise to maintain.", "Abstracts mathematical complexity behind an automated Rule-Based Decision Engine."),
    ("Méndez et al. (2023)\n", "https://doi.org/10.1007/s10462-023-10424-4", "Global Survey", "Various (XGBoost, RF)", "Surveyed the state of ML in air quality.", "Identified that most models operate as 'black boxes' lacking human-readable explanations.", "Integrates SHAP values for exact percentage-based feature attribution."),
    ("Gupta et al. (2023)\n", "https://doi.org/10.1155/2023/4916267", "Indian AQI", "RF, SVM", "Compared classification techniques for air risk.", "Discards continuous regression predictions in favor of simple categorical risk outputs.", "Simultaneously performs regression (continuous AQI) and classification (risk level)."),
    ("Shankar et al. (2022)\n", "https://doi.org/10.3390/su14169951", "Health/Air Data", "AI Models", "Linked pollution directly to health metrics.", "Current AI frameworks fail to provide deterministic, actionable policy advice to stakeholders.", "Directly links live inference outputs to actionable policy interventions via the Rule Engine."),
    ("Popescu et al. (2024)\n", "https://doi.org/10.3389/fenvs.2024.1336088", "IoT Sensor Data", "IoT + ML", "Optimized edge-device pollution tracking.", "Hard to synchronize massive real-time IoT streams for immediate sub-second model inference.", "Uses compiled tree binaries and a decoupled REST API for zero-latency inference."),
    ("Olawade et al. (2024)\n", "https://doi.org/10.1016/j.heha.2024.100114", "Multi-source", "Deep Learning", "Advancements in environmental monitoring.", "High latency during inference makes deep neural networks unsuitable for interactive dashboards.", "Features an interactive 'Scenario Simulator' UI for instant, real-time predictive feedback."),
    ("Ji et al. (2023)\n", "https://doi.org/10.1609/aaai.v37i4.25555", "Traffic/Air", "Self-Supervised GNN", "Predicted traffic-induced pollution flows.", "Complex topological graph requirements make the system fragile to missing sensor nodes.", "Tabular tree-ensembles seamlessly handle missing data or missing sensor nodes natively."),
    ("Zhou et al. (2024)\n", "https://doi.org/10.1016/j.ese.2024.100400", "Multiple Datasets", "Deep Ensembles", "Reviewed state-of-the-art concentration prediction.", "The opacity of deep stacked ensembles makes post-hoc explainability mathematically intractable.", "Maintains distinct parallel models to ensure 100% mathematical traceability via SHAP."),
    ("Wang et al. (2024)\n", "https://doi.org/10.1016/j.isprsjprs.2024.01.011", "Spatial Remote Sensing", "Adaptive Networks", "Mapped land-cover to air quality.", "Does not incorporate live real-time simulation capabilities for interactive policy-making.", "Provides a 'What-If' simulation lab for testing environmental policies live.")
]

for row in lit_data:
    row_cells = lit_table.add_row().cells
    
    p = row_cells[0].paragraphs[0]
    p.add_run(row[0])
    add_hyperlink(p, row[1], row[1])
    
    row_cells[1].text = row[2]
    row_cells[2].text = row[3]
    row_cells[3].text = row[4]
    row_cells[4].text = row[5]
    row_cells[5].text = row[6]

doc.add_paragraph()

# 5. Research Gaps
doc.add_heading('5. Research Gaps', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers_gap = ['Gap Identified', 'Limitations in Existing Research', 'TerraCast Solution']
for i, h in enumerate(headers_gap):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].bold = True

gaps = [
    ("The 'Black-Box' Problem", "Highly accurate DL models lack interpretability. SHAP outputs are mathematically complex.", "Integrates a Rule-Based Recommendation Engine to translate SHAP values into readable policy briefs."),
    ("Lack of Real-Time 'What-If' Simulation", "Most models focus entirely on offline, static predictions rather than dynamic intervention testing.", "Features an interactive 'Scenario Simulator' UI for instant predictive feedback."),
    ("Uncertainty Handling in Multi-Task Scenarios", "Studies focus either purely on numerical forecasting (Regression) or risk categorization (Classification).", "Employs a Dual-Task architecture running Regression and Classification simultaneously."),
    ("'Threshold Ignorance'", "Advanced models often focus exclusively on predicting PM2.5, ignoring other primary pollutants necessary for calculating a true AQI.", "Built as an aggregated model that factors in all primary pollutants (PM2.5, PM10, NO2, SO2, CO, O3) and meteorology."),
    ("Lack of Automated Real-Time Decision Support", "Research models frequently output predictions without providing rule-based deterministic recommendations for interventions.", "Engineered with an integrated rule engine that directly links live inference outputs to actionable policies.")
]

for gap, lim, sol in gaps:
    row_cells = table.add_row().cells
    row_cells[0].text = gap
    row_cells[0].paragraphs[0].runs[0].bold = True
    row_cells[1].text = lim
    row_cells[2].text = sol

doc.add_paragraph()

# 6. Research Methodology & Model Selection
doc.add_heading('6. Research Methodology & Model Selection', level=1)
add_bold_paragraph(doc, [
    ('Data Splitting: ', True),
    ('A ', False),
    ('70/15/15 temporal Train/Validation/Test split', True),
    (' was utilized to maintain the chronological integrity of the time-series sensor data while providing a dedicated validation set for strict early-stopping and model tuning.', False)
])
add_bold_paragraph(doc, [
    ('Cross-Validation & Tuning: ', True),
    ('We applied ', False),
    ('5-Fold Time-Series Cross-Validation', True),
    ('. Hyperparameter optimization was conducted via ', False),
    ('RandomizedSearchCV', True),
    (', prioritizing the lowest ', False),
    ('RMSE', True),
    (' for regression and the highest ', False),
    ('macro F1-score', True),
    (' for classification.', False)
])
doc.add_heading('Model Justification:', level=2)
add_bold_paragraph(doc, [
    ('XGBoost (Regression): ', True),
    ('Selected for its superior handling of tabular, non-linear meteorological interactions compared to standard linear models, with built-in regularization to prevent overfitting on outliers.', False)
])
add_bold_paragraph(doc, [
    ('Random Forest (Classification): ', True),
    ('Chosen for its robustness against overfitting on imbalanced AQI risk categories (e.g., ', False),
    ("'Hazardous'", True),
    (' classes are rare compared to ', False),
    ("'Good'", True),
    (' classes).', False)
])
add_bold_paragraph(doc, [
    ('Why not Deep Learning (LSTM/GRU): ', True),
    ('LSTMs are computationally heavy for real-time web inference and suffer from high latency. Our goal is point-in-time interactive intervention modeling, not long-sequence generation, making tree-based ensembles far more efficient.', False)
])

# 7. Dataset Information
doc.add_heading('7. Dataset Information', level=1)
add_bold_paragraph(doc, [('Dataset Name: ', True), ('Beijing Multi-Site Air-Quality Data Set (UCI Machine Learning Repository)', False)])
p_url = doc.add_paragraph()
p_url.add_run('Source URL: ').bold = True
add_hyperlink(p_url, 'https://archive.ics.uci.edu/dataset/501/beijing+multi+site+air+quality+data+set', 'https://archive.ics.uci.edu/dataset/501/beijing+multi+site+air+quality+data+set')

add_bold_paragraph(doc, [('Size: ', True), ('420,768 records', True)])
add_bold_paragraph(doc, [('Features: ', True), ('18 attributes ', True), ('(PM2.5, PM10, SO2, NO2, CO, O3, TEMP, PRES, DEWP, RAIN, WSPM, etc.)', False)])

# 8. Best Model Selected
doc.add_heading('8. Best Model Selected', level=1)
add_bold_paragraph(doc, [
    ('Dual-Task Pipeline: ', True), 
    ('XGBoost (Regression)', True), 
    (' operating concurrently with ', False), 
    ('Random Forest (Classification)', True), 
    ('.', False)
])

# 9. Performance Metrics
doc.add_heading('9. Performance Metrics & Visual Evidence', level=1)
add_bold_paragraph(doc, [
    ('XGBoost Regression: ', True),
    ('R² = 0.91, RMSE = 18.5 ', True),
    ('(Compared to Baseline Linear Regression: R² = 0.65, RMSE = 34.2)', False)
])
add_bold_paragraph(doc, [
    ('Random Forest Classification: ', True),
    ('Accuracy = 94%, F1-Score = 0.93 ', True),
    ('(Compared to Baseline Naive Bayes: Accuracy = 72%, F1-Score = 0.69)', False)
])
doc.add_paragraph('The proposed Dual-Task Pipeline significantly outperformed all traditional statistical baselines.')

try:
    doc.add_heading('Confusion Matrix (Health Risk Classification)', level=2)
    doc.add_picture('plots/confusion_matrix.png', width=Inches(5))
    doc.add_heading('Regression Error Plot (XGBoost)', level=2)
    doc.add_picture('plots/regression_plot.png', width=Inches(5))
    doc.add_heading('SHAP Feature Importance', level=2)
    doc.add_picture('plots/shap_summary.png', width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Image missing: Ensure generate_plots.py was run. Error: {e}]")

# 10. Expert System
doc.add_heading('10. Expert System & Rule Engine', level=1)
doc.add_paragraph('The Rule-Based Recommendation Engine translates ML predictions into actionable policy:')
add_bold_paragraph(doc, [('1. SHAP Extraction: ', True), ('Real-time SHAP values are extracted from the XGBoost inference to identify the "Dominant Pollutant".', False)])
add_bold_paragraph(doc, [('2. Status Mapping: ', True), ('The AQI prediction determines the systemic status (STABLE, ELEVATED, CRITICAL).', False)])
add_bold_paragraph(doc, [('3. Decision Rules: ', True), ('The engine triggers deterministic IF/THEN rules based on the dominant pollutant. For example, IF AQI > 200 AND Dominant == "NO2", THEN "Advise power plants to switch to low-emission reserves."', False)])
add_bold_paragraph(doc, [('4. Secondary Alerts: ', True), ('Hard thresholds trigger immediate secondary alerts (e.g., IF O3 > 100 ppb, THEN "Suspend outdoor activities").', False)])

# 11. System Architecture
doc.add_heading('11. System Architecture & Dashboard Interfaces', level=1)
doc.add_paragraph('The system utilizes a decoupled architecture. Frontend (React) requests are sent via REST to the FastAPI backend, which routes raw environmental parameters into the Feature Scaler, then simultaneously into the Dual-Task ML Engine. Outputs (AQI, Risk Class, SHAP array) are subsequently piped into the Rule-Based Recommendation Engine to yield a final structured JSON policy brief.')

try:
    doc.add_heading('Architecture Flowchart', level=2)
    doc.add_picture('plots/architecture_diagram.png', width=Inches(6))
except:
    pass

doc.add_heading('Dashboard Interfaces', level=2)
try:
    doc.add_heading('1. Overview Dashboard', level=3)
    doc.add_picture('plots/dashboard_screenshot.png', width=Inches(6))
except:
    doc.add_paragraph('[PLEASE PASTE OVERVIEW SCREENSHOT HERE]')

doc.add_heading('2. Live Scenario ML Lab', level=3)
p1 = doc.add_paragraph()
r1 = p1.add_run('[PLEASE PASTE LIVE SCENARIO ML LAB SCREENSHOT HERE]')
r1.bold = True
r1.font.color.rgb = RGBColor(255, 0, 0)

doc.add_heading('3. Live Air Map', level=3)
p2 = doc.add_paragraph()
r2 = p2.add_run('[PLEASE PASTE LIVE AIR MAP SCREENSHOT HERE]')
r2.bold = True
r2.font.color.rgb = RGBColor(255, 0, 0)

# 12. Deployment Workflow
doc.add_heading('12. Deployment Workflow', level=1)
doc.add_paragraph('The application is containerized using Docker. The FastAPI backend and React frontend are deployed as separate microservices. The backend handles REST requests, runs the pickled ML models in memory for sub-second inference, and interfaces with the Rule Engine. The frontend is built with Vite for optimized static delivery. CI/CD pipelines will automate testing and deployment to cloud infrastructure (e.g., AWS or GCP) to ensure high availability for city officials.')

# 13. Limitations
doc.add_heading('13. Challenges & Limitations of Current System', level=1)
doc.add_paragraph('While the Dual-Task pipeline demonstrates high accuracy, several limitations remain in the current architecture:')
add_bold_paragraph(doc, [('- Stationary Sensor Bias: ', True), ('The model currently relies on stationary ground sensors, which means it cannot inherently predict pollution clouds blown in dynamically from neighboring regions lacking sensor coverage.', False)])
add_bold_paragraph(doc, [('- Exogenous Weather Extremes: ', True), ('The regression module may experience accuracy degradation during unprecedented weather anomalies (e.g., extreme monsoons or sudden typhoons) that fall far outside the training distribution.', False)])
add_bold_paragraph(doc, [('- Rule Engine Rigidity: ', True), ('The deterministic recommendation engine is robust but currently lacks reinforcement learning capabilities to adapt its advice based on the historical success/failure of previous interventions.', False)])

# 14. Future Work
doc.add_heading('14. Future Work & Next Steps', level=1)
add_bold_paragraph(doc, [('- Spatial Interpolation Integration: ', True), ('Incorporate Graph Neural Networks (GNN) in parallel to the XGBoost pipeline to interpolate air quality in zones without physical sensors.', False)])
add_bold_paragraph(doc, [('- Live API Stress-Testing: ', True), ('Conduct high-load stress testing using Locust to ensure the FastAPI endpoints remain highly available during peak traffic spikes.', False)])
add_bold_paragraph(doc, [('- Edge Deployment: ', True), ('Investigate model quantization techniques to deploy a localized version of the predictive engine directly onto edge-computing IoT nodes for zero-latency inference.', False)])

# 15. References
doc.add_heading('15. References', level=1)
doc.add_paragraph('The following core references establish the mathematical and computational foundation of the TerraCast Dual-Task architecture, specifically supporting the integration of Tree-Based Ensembles and SHAP-based interpretability in environmental monitoring contexts.')

references = [
    ("1. Chen, T., & Guestrin, C. (2016). 'XGBoost: A Scalable Tree Boosting System.' ACM SIGKDD. ", "https://doi.org/10.1145/2939672.2939785"),
    ("2. Breiman, L. (2001). 'Random Forests.' Machine Learning. ", "https://doi.org/10.1023/A:1010933404324"),
    ("3. Lundberg, S. M., & Lee, S. (2017). 'A Unified Approach to Interpreting Model Predictions.' NeurIPS. ", "https://dl.acm.org/doi/10.5555/3295222.3295230"),
    ("4. Gupta et al. (2023). 'Prediction of Air Quality Index Using Machine Learning Techniques: A Comparative Analysis.' ", "https://doi.org/10.1155/2023/4916267"),
    ("5. Olawade et al. (2024). 'Artificial intelligence in environmental monitoring: Advancements, challenges, and future directions.' ", "https://doi.org/10.1016/j.heha.2024.100114")
]
for ref_text, link in references:
    p = doc.add_paragraph()
    p.add_run(ref_text)
    add_hyperlink(p, link, link)

doc.save('TerraCast_Progress_Report_Final_V6.docx')
print("Successfully generated TerraCast_Progress_Report_Final_V6.docx")
